import os
from pymongo import MongoClient
from datetime import datetime
from email.message import EmailMessage
import smtplib
from docx import Document
from io import BytesIO
import sys

def to_12hr(dt_obj):
   
    if dt_obj in ("-", None):
        return "-"
    
    if isinstance(dt_obj, datetime):
        return dt_obj.strftime("%I:%M %p")
    
    try:
        # Fallback for old string-based data if any remains
        t = datetime.strptime(str(dt_obj), "%H:%M")
        return t.strftime("%I:%M %p")
    except Exception:
        return str(dt_obj)

try:
    MONGO_URI = os.environ["MONGO_URI"]
    GMAIL_USER = os.environ["GMAIL_USER"]
    GMAIL_APP_PASSWORD = os.environ["GMAIL_APP_PASSWORD"]
    RECIPIENT_EMAILS = [e.strip() for e in os.environ["RECIPIENT_EMAILS"].split(",")]
except KeyError as e:
    print(f"Missing required environment variable: {e}")
    sys.exit(1)

DB_NAME = os.environ.get("DB_NAME", "infoDB")
COLLECTION_NAME = os.environ.get("COLLECTION_NAME", "Log_In")

try:
    client = MongoClient(MONGO_URI)
    db = client[DB_NAME]
    col = db[COLLECTION_NAME]
except Exception as e:
    print("MongoDB connection failed:", e)
    sys.exit(1)


doc = col.find_one(sort=[("_id", -1)])

if not doc or not doc.get("logs"):
    print("No logs found. Exiting.")
    sys.exit(0)

logs = doc["logs"]
week_start = doc.get("week_start", "Unknown")
week_end = doc.get("week_end", "Unknown")

lines = []
lines.append("Weekly Attendance Report")
lines.append(f"Week: {week_start} -> {week_end}")
lines.append("-" * 30)

for day in sorted(logs.keys()):
    lines.append(f"\nDATE: {day}")
    users = logs[day]

    for user, sessions in users.items():
       
        if isinstance(sessions, list):
            for i, session in enumerate(sessions, 1):
                sign_in = to_12hr(session.get("sign_in", "-"))
                sign_out = to_12hr(session.get("sign_out", "-"))
                lines.append(f"  {user} (Session {i}): {sign_in} -> {sign_out}")
        else:
            # Fallback for old data structure
            sign_in = to_12hr(sessions.get("sign_in", "-"))
            sign_out = to_12hr(sessions.get("sign_out", "-"))
            lines.append(f"  {user}: {sign_in} -> {sign_out}")

text = "\n".join(lines)

doc_buffer = BytesIO()
document = Document()
document.add_heading('Weekly Attendance Report', 0)
document.add_paragraph(f"Week: {week_start} to {week_end}")

for line in lines[3:]: 
    document.add_paragraph(line)

document.save(doc_buffer)
doc_buffer.seek(0)

msg = EmailMessage()
msg["Subject"] = f"Weekly Attendance Report ({week_start})"
msg["From"] = GMAIL_USER
msg["To"] = ", ".join(RECIPIENT_EMAILS)

msg.set_content(text)

msg.add_attachment(
    doc_buffer.read(),
    maintype="application",
    subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
    filename=f"weekly_report_{week_start}.docx"
)

try:
    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(GMAIL_USER, GMAIL_APP_PASSWORD)
        server.send_message(msg)
    print("Email sent successfully.")
except Exception as e:
    print("Failed to send email:", e)
    sys.exit(1)
